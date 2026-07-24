"""Turn a resume file (PDF / DOCX / TXT / MD) into plain text — as complete as
possible so the LLM can extract everything.

PDF strategy: try pdfplumber (layout-aware, keeps more of the real text), and
fall back to pypdf if pdfplumber isn't available or yields nothing. If both fail
to produce text, the PDF is almost certainly a scanned image and we say so
clearly rather than silently returning garbage.
"""

from __future__ import annotations

from pathlib import Path


def extract_text(path: str | Path) -> str:
    p = Path(path).expanduser()
    if not p.exists():
        raise FileNotFoundError(f"Resume not found: {p}")

    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return _pdf(p)
    if suffix == ".docx":
        return _docx(p)
    if suffix in {".txt", ".md", ".text"}:
        return p.read_text(errors="ignore")
    # Best effort: try to read as text.
    try:
        return p.read_text(errors="ignore")
    except Exception as e:  # noqa: BLE001
        raise ValueError(f"Unsupported resume format: {suffix}") from e


def _pdf(p: Path) -> str:
    text = _pdf_pdfplumber(p) or _pdf_pypdf(p)
    text = (text or "").strip()
    if not text:
        raise ValueError(
            "Could not extract any text from this PDF — it looks like a scanned "
            "image (no selectable text). Please provide a text-based PDF/DOCX, or "
            "paste your resume text into a .txt file."
        )
    return text


def _pdf_pdfplumber(p: Path) -> str:
    try:
        import pdfplumber
    except Exception:  # noqa: BLE001 — not installed
        return ""
    try:
        parts: list[str] = []
        with pdfplumber.open(str(p)) as pdf:
            for page in pdf.pages:
                parts.append(page.extract_text(x_tolerance=1.5, y_tolerance=3) or "")
        return "\n".join(parts)
    except Exception:  # noqa: BLE001 — fall back to pypdf
        return ""


def _pdf_pypdf(p: Path) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(p))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception:  # noqa: BLE001
        return ""


def _docx(p: Path) -> str:
    import docx

    doc = docx.Document(str(p))
    # Paragraphs + table cells, so nothing in a two-column resume layout is lost.
    parts = [para.text for para in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts).strip()
